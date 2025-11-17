from __future__ import annotations

import re
from io import BytesIO
from typing import Optional, Tuple

from docx import Document as DocxDocument
from fastapi import UploadFile
from langdetect import detect, LangDetectException
from PyPDF2 import PdfReader

from app.models import DocumentMetadata, ParsedDocument


async def parse_document(file: UploadFile) -> ParsedDocument:
    raw_bytes = await file.read()
    file_ext = (file.filename or "").lower()

    if file_ext.endswith(".pdf"):
        text, pages = _read_pdf(raw_bytes)
    elif file_ext.endswith(".docx"):
        text, pages = _read_docx(raw_bytes)
    elif file_ext.endswith(".txt"):
        text = raw_bytes.decode("utf-8", errors="ignore")
        pages = None
    else:
        text = raw_bytes.decode("utf-8", errors="ignore")
        pages = None

    word_count = len(re.findall(r"\b\w+\b", text))
    language = _detect_language(text)

    metadata = DocumentMetadata(
        filename=file.filename or "untitled",
        file_type=file_ext.split(".")[-1] if "." in file_ext else "unknown",
        pages=pages,
        word_count=word_count,
        language=language,
    )

    return ParsedDocument(
        metadata=metadata,
        text=text,
    )


def _read_pdf(raw_bytes: bytes) -> Tuple[str, Optional[int]]:
    reader = PdfReader(BytesIO(raw_bytes))
    pages = len(reader.pages)
    text_chunks = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(text_chunks), pages


def _read_docx(raw_bytes: bytes) -> Tuple[str, Optional[int]]:
    document = DocxDocument(BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text]
    return "\n".join(paragraphs), None


def _detect_language(text: str) -> Optional[str]:
    try:
        return detect(text) if text.strip() else None
    except LangDetectException:
        return None

